from __future__ import annotations

from pathlib import Path

from docx import Document

from md2word.pipeline import BuildProfile, _apply_captions, _validate_output, captions_from_markdown


def test_caption_parser_rejects_duplicate_labels() -> None:
    profile = BuildProfile(reference_docx=Path("template.docx"))
    source = "表2-1 First\n\n表2-1 First\n"
    try:
        captions_from_markdown(source, profile)
    except ValueError as error:
        assert "unique" in str(error)
    else:
        raise AssertionError("Expected duplicate caption validation to fail")


def test_caption_fields_do_not_use_chapter_sequences(tmp_path: Path) -> None:
    output = tmp_path / "sample.docx"
    document = Document()
    document.add_paragraph("表2-1 Verification")
    document.add_paragraph("图2-2 Delivery flow")
    document.save(output)

    profile = BuildProfile(reference_docx=output)
    captions = captions_from_markdown("表2-1 Verification\n图2-2 Delivery flow\n", profile)
    document = Document(output)
    _apply_captions(document, captions, profile)
    document.save(output)

    _validate_output(output, captions, [], profile)
