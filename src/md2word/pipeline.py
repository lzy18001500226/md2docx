from __future__ import annotations

import re
import shutil
import subprocess
import tempfile
import tomllib
import zipfile
from dataclasses import dataclass
from hashlib import sha256
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree


IMAGE_RE = re.compile(r"!\[[^\]]*\]\((?P<path>[^)\s]+)(?:\s+[^)]*)?\)")


@dataclass(frozen=True)
class BuildProfile:
    reference_docx: Path
    figure_label: str = "图"
    table_label: str = "表"
    code_style: str = "Normal"


@dataclass(frozen=True)
class Caption:
    kind: str
    chapter: int
    sequence: int
    title: str

    @property
    def visible_text(self) -> str:
        return f"{self.kind}{self.chapter}-{self.sequence} {self.title}"


@dataclass(frozen=True)
class BuildResult:
    output: Path
    caption_count: int
    image_count: int


def load_profile(path: Path) -> BuildProfile:
    path = path.resolve()
    values = tomllib.loads(path.read_text(encoding="utf-8"))
    reference = (path.parent / values["reference_docx"]).resolve()
    if not reference.is_file():
        raise FileNotFoundError(f"Profile reference DOCX does not exist: {reference}")
    return BuildProfile(
        reference_docx=reference,
        figure_label=str(values.get("figure_label", "图")),
        table_label=str(values.get("table_label", "表")),
        code_style=str(values.get("code_style", "Normal")),
    )


def _caption_pattern(profile: BuildProfile) -> re.Pattern[str]:
    labels = "|".join(re.escape(value) for value in (profile.figure_label, profile.table_label))
    return re.compile(
        rf"^(?P<kind>{labels})\s*(?P<chapter>\d+)\s*-\s*"
        rf"(?P<sequence>\d+)\s+(?P<title>.+?)\s*$"
    )


def captions_from_markdown(markdown: str, profile: BuildProfile) -> list[Caption]:
    pattern = _caption_pattern(profile)
    captions: list[Caption] = []
    for line in markdown.splitlines():
        match = pattern.match(line.strip())
        if match:
            captions.append(
                Caption(
                    kind=match.group("kind"),
                    chapter=int(match.group("chapter")),
                    sequence=int(match.group("sequence")),
                    title=match.group("title"),
                )
            )
    if len({caption.visible_text for caption in captions}) != len(captions):
        raise ValueError("Caption labels must be unique.")
    return captions


def _markdown_image_paths(source: Path, markdown: str) -> list[Path]:
    paths: list[Path] = []
    for match in IMAGE_RE.finditer(markdown):
        raw_path = match.group("path")
        if raw_path.startswith(("http://", "https://", "data:")):
            continue
        image = (source.parent / raw_path).resolve()
        if not image.is_file():
            raise FileNotFoundError(f"Markdown image does not exist: {raw_path}")
        paths.append(image)
    return paths


def _clear_paragraph(paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def _append_field(paragraph, instruction: str, visible_value: str) -> None:
    begin = paragraph.add_run()
    begin_element = OxmlElement("w:fldChar")
    begin_element.set(qn("w:fldCharType"), "begin")
    begin._r.append(begin_element)

    instruction_run = paragraph.add_run()
    instruction_element = OxmlElement("w:instrText")
    instruction_element.set(qn("xml:space"), "preserve")
    instruction_element.text = f" {instruction} "
    instruction_run._r.append(instruction_element)

    separator = paragraph.add_run()
    separator_element = OxmlElement("w:fldChar")
    separator_element.set(qn("w:fldCharType"), "separate")
    separator._r.append(separator_element)

    paragraph.add_run(visible_value)

    end = paragraph.add_run()
    end_element = OxmlElement("w:fldChar")
    end_element.set(qn("w:fldCharType"), "end")
    end._r.append(end_element)


def _apply_captions(document: Document, captions: list[Caption], profile: BuildProfile) -> None:
    by_text = {caption.visible_text: caption for caption in captions}
    for paragraph in document.paragraphs:
        caption = by_text.get(paragraph.text.strip())
        if caption is None:
            continue
        _clear_paragraph(paragraph)
        if "Caption" in [style.name for style in document.styles]:
            paragraph.style = "Caption"
        paragraph.add_run(f"{caption.kind}{caption.chapter}-")
        field_name = "Figure" if caption.kind == profile.figure_label else "Table"
        _append_field(paragraph, f"SEQ {field_name} \\r {caption.sequence} \\* ARABIC", str(caption.sequence))
        paragraph.add_run(f" {caption.title}")


def _normalize_code_styles(document: Document, code_style: str) -> None:
    available = {style.name for style in document.styles}
    if code_style not in available:
        raise ValueError(f"Configured code style is absent from template: {code_style}")
    for paragraph in document.paragraphs:
        if paragraph.style.name == "Source Code":
            paragraph.style = code_style


def _set_document_settings(output: Path) -> None:
    namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    settings_path = "word/settings.xml"
    with zipfile.ZipFile(output, "r") as archive:
        files = {info.filename: archive.read(info.filename) for info in archive.infolist()}
    settings = etree.fromstring(files[settings_path])
    for name, value in (("doNotCompressPictures", None), ("updateFields", "true")):
        node = settings.find(f"{{{namespace}}}{name}")
        if node is None:
            node = etree.Element(f"{{{namespace}}}{name}")
            settings.append(node)
        if value is not None:
            node.set(f"{{{namespace}}}val", value)
    files[settings_path] = etree.tostring(
        settings, xml_declaration=True, encoding="UTF-8", standalone=True
    )
    temporary = output.with_suffix(".docx.tmp")
    with zipfile.ZipFile(temporary, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for name, content in files.items():
            archive.writestr(name, content)
    temporary.replace(output)


def _update_with_word(output: Path, pdf_output: Path | None) -> None:
    output = output.resolve()
    pdf_output = pdf_output.resolve() if pdf_output else None
    ps_quote = lambda value: "'" + str(value).replace("'", "''") + "'"
    export_pdf = "$true" if pdf_output else "$false"
    pdf_literal = ps_quote(pdf_output) if pdf_output else "$null"
    script = f"""
$ErrorActionPreference = 'Stop'
$word = New-Object -ComObject Word.Application
$word.Visible = $false
$word.DisplayAlerts = 0
try {{
  $document = $word.Documents.Open({ps_quote(output)})
  $document.Fields.Update() | Out-Null
  $document.Save()
  if ({export_pdf}) {{ $document.ExportAsFixedFormat({pdf_literal}, 17) }}
  $document.Close()
}} finally {{
  $word.Quit()
}}
"""
    subprocess.run(["powershell.exe", "-NoProfile", "-NonInteractive", "-Command", script], check=True)


def _validate_output(
    output: Path,
    captions: list[Caption],
    source_images: list[Path],
    profile: BuildProfile,
) -> None:
    with zipfile.ZipFile(output) as archive:
        invalid = archive.testzip()
        if invalid:
            raise RuntimeError(f"DOCX ZIP integrity failed at {invalid}")
        document = etree.fromstring(archive.read("word/document.xml"))
        embedded_hashes = {
            sha256(archive.read(info.filename)).hexdigest()
            for info in archive.infolist()
            if info.filename.startswith("word/media/")
        }
    namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    fields = [" ".join(value.split()) for value in document.xpath(".//w:instrText/text()", namespaces={"w": namespace})]
    if any(field.startswith("SEQ Chapter") for field in fields):
        raise RuntimeError("Output retains SEQ Chapter fields.")
    for caption in captions:
        field_name = "Figure" if caption.kind == profile.figure_label else "Table"
        expected = f"SEQ {field_name} \\r {caption.sequence} \\* ARABIC"
        if expected not in fields:
            raise RuntimeError(f"Caption field missing: {caption.visible_text}")
    source_hashes = {sha256(image.read_bytes()).hexdigest() for image in source_images}
    missing = source_hashes - embedded_hashes
    if missing:
        raise RuntimeError(f"Output omitted {len(missing)} source image(s).")


def _run_pandoc(source: Path, reference_docx: Path, output: Path) -> None:
    pandoc = shutil.which("pandoc")
    if pandoc is None:
        raise RuntimeError("Pandoc was not found on PATH.")
    lines = source.read_text(encoding="utf-8").splitlines(keepends=True)
    title_index = next((index for index, line in enumerate(lines) if line.startswith("# ")), None)
    title = lines[title_index][2:].strip() if title_index is not None else source.stem
    if title_index is not None:
        del lines[title_index]
        if title_index < len(lines) and not lines[title_index].strip():
            del lines[title_index]

    temporary_path: Path | None = None
    try:
        with tempfile.NamedTemporaryFile(
            mode="w", encoding="utf-8", suffix=".md", dir=source.parent, delete=False
        ) as temporary:
            temporary.write("".join(lines))
            temporary_path = Path(temporary.name)
        subprocess.run(
            [
                pandoc,
                str(temporary_path),
                "--from=markdown+tex_math_dollars+tex_math_single_backslash",
                "--to=docx",
                "--standalone",
                "--reference-doc",
                str(reference_docx),
                "--resource-path",
                str(source.parent.resolve()),
                "--metadata",
                f"title={title}",
                "--output",
                str(output.resolve()),
            ],
            cwd=source.parent,
            check=True,
        )
    finally:
        if temporary_path is not None:
            temporary_path.unlink(missing_ok=True)


def build_document(
    source: Path,
    profile: BuildProfile,
    output: Path,
    *,
    overwrite: bool = False,
    use_word: bool = False,
    pdf_output: Path | None = None,
) -> BuildResult:
    source = source.resolve()
    output = output.resolve()
    if not source.is_file():
        raise FileNotFoundError(source)
    if output.exists() and not overwrite:
        raise FileExistsError(f"Refusing to overwrite output: {output}")
    output.parent.mkdir(parents=True, exist_ok=True)
    markdown = source.read_text(encoding="utf-8")
    captions = captions_from_markdown(markdown, profile)
    images = _markdown_image_paths(source, markdown)
    _run_pandoc(source, profile.reference_docx, output)
    document = Document(output)
    _apply_captions(document, captions, profile)
    _normalize_code_styles(document, profile.code_style)
    document.save(output)
    _set_document_settings(output)
    if use_word:
        _update_with_word(output, pdf_output)
    _validate_output(output, captions, images, profile)
    return BuildResult(output=output, caption_count=len(captions), image_count=len(images))
